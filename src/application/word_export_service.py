"""
Word document export service.
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any
from pathlib import Path


class WordExportService:
    """Word 문서 생성 서비스."""
    
    def create_notice_document(self, notice_data: Dict[str, Any], output_path: str) -> str:
        """공지사항을 Word 문서로 생성.
        
        Args:
            notice_data: 공지사항 데이터
            output_path: 출력 파일 경로
            
        Returns:
            생성된 파일 경로
        """
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading(notice_data['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 구분선
        doc.add_paragraph("=" * 50)
        
        # 메타데이터 테이블
        doc.add_heading('📋 공지사항 정보', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # 테이블 내용
        info_data = [
            ('공지번호', notice_data['id']),
            ('카테고리', notice_data['category']),
            ('작성일', notice_data['created_date'][:10] if notice_data['created_date'] else 'N/A'),
            ('조회수', str(notice_data['view_count'])),
            ('첨부파일', '있음' if notice_data.get('has_attachment') else '없음')
        ]
        
        for i, (key, value) in enumerate(info_data):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = value
        
        # 구분선
        doc.add_paragraph("=" * 50)
        
        # 본문 내용
        if notice_data.get('content'):
            doc.add_heading('📄 내용', level=1)
            content_para = doc.add_paragraph(notice_data['content'])
        else:
            doc.add_heading('📄 내용', level=1)
            doc.add_paragraph('상세 내용이 없습니다.')
        
        # 구분선
        doc.add_paragraph("=" * 50)
        
        # 생성 정보
        footer = doc.add_paragraph()
        footer.add_run(f"📅 문서 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        footer.add_run("🌐 출처: TOPIS 서울시 교통정보센터\n")
        footer.add_run("🔗 URL: https://topis.seoul.go.kr")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 디렉토리 생성
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # 파일 저장
        doc.save(output_path)
        return output_path
    
    def create_multiple_notices_document(self, notices_data: list, output_path: str) -> str:
        """여러 공지사항을 하나의 Word 문서로 생성.
        
        Args:
            notices_data: 공지사항 데이터 리스트
            output_path: 출력 파일 경로
            
        Returns:
            생성된 파일 경로
        """
        doc = Document()
        
        # 표지
        cover_title = doc.add_heading('TOPIS 공지사항 모음', 0)
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"📊 총 {len(notices_data)}개의 공지사항")
        doc.add_paragraph(f"📅 생성일: {datetime.now().strftime('%Y-%m-%d')}")
        
        # 페이지 나누기
        doc.add_page_break()
        
        # 각 공지사항 추가
        for i, notice in enumerate(notices_data, 1):
            # 공지사항 번호
            doc.add_heading(f'{i}. {notice["title"]}', level=1)
            
            # 메타데이터
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"🆔 공지번호: {notice['id']} | ")
            meta_para.add_run(f"📂 카테고리: {notice['category']} | ")
            meta_para.add_run(f"📅 작성일: {notice['created_date'][:10] if notice['created_date'] else 'N/A'} | ")
            meta_para.add_run(f"👀 조회수: {notice['view_count']}")
            
            # 내용
            if notice.get('content'):
                doc.add_paragraph(notice['content'])
            else:
                doc.add_paragraph('상세 내용이 없습니다.')
            
            # 구분선 (마지막이 아닌 경우)
            if i < len(notices_data):
                doc.add_paragraph("─" * 50)
        
        # 디렉토리 생성
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # 파일 저장
        doc.save(output_path)
        return output_path
